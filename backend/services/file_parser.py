"""
File parser service — extracts text from PDF, DOCX, and TXT files.
"""

import io
from PyPDF2 import PdfReader
from docx import Document


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts).strip()


def parse_txt(file_bytes: bytes) -> str:
    """Extract text from a plain text file."""
    # Try UTF-8 first, fall back to latin-1
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1").strip()


def parse_file(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file format and extract text.

    Args:
        file_bytes: Raw file content
        filename: Original filename (used to detect format)

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format is not supported
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return parse_docx(file_bytes)
    elif ext == "txt":
        return parse_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: .{ext}. "
            f"Supported formats: PDF, DOCX, DOC, TXT"
        )
